"""Generate the JBS Ethics Admin user manual as a Word document."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "ethics_admin_manual_assets"
OUTPUT = ROOT / "docs" / "JBS_Ethics_Admin_User_Manual.docx"
LOGO = ROOT / "app" / "ethics_production_app" / "static" / "img" / "image.png"

ORANGE = "EF820D"
DARK_BLUE = "022169"
LIGHT_ORANGE = "FFF2E5"
LIGHT_GREY = "F2F3F5"
WHITE = "FFFFFF"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), color)
    tc_pr.append(fill)


def set_cell_text_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char1, instr_text, fld_char2))


def add_callout(doc, title, text, color=LIGHT_ORANGE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    bold = p.add_run(f"{title}: ")
    bold.bold = True
    bold.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_steps(doc, steps):
    for text in steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(text)


def add_bullets(doc, items):
    for text in items:
        doc.add_paragraph(text, style="List Bullet")


def add_screenshot(doc, filename, caption):
    path = ASSETS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.65))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles["Caption"]


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in (
        ("Title", 28, DARK_BLUE),
        ("Heading 1", 18, DARK_BLUE),
        ("Heading 2", 14, ORANGE),
        ("Heading 3", 11, DARK_BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.add_run("JBS Ethics Application — Admin User Manual | ")
        add_page_number(footer)


def add_cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.5))
    title = doc.add_paragraph("JBS Ethics Application", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("ADMIN USER MANUAL")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(21)
    subtitle.runs[0].font.bold = True
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(ORANGE)
    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    values = (
        ("System", "Integrated Ethics and MBA — Ethics module"),
        ("Audience", "Ethics Administrators and Super Administrators"),
        ("Version", "1.0"),
        ("Updated", "11 August 2026"),
    )
    for row, (label, value) in zip(meta.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        shade(row.cells[0], ORANGE)
        set_cell_text_color(row.cells[0], WHITE)
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()
    add_callout(
        doc,
        "Purpose",
        "This manual explains the daily Admin workflow, from managing users and submissions to REC processing, certificates, reporting, and audit monitoring.",
    )
    doc.add_paragraph("Internal operational document", style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_contents(doc):
    doc.add_heading("Contents", level=1)
    sections = (
        "1. Access and navigation",
        "2. Admin dashboard and submission folders",
        "3. Reviewing and routing an application",
        "4. User management",
        "5. Supervisors, reviewers and password resets",
        "6. Archived student forms",
        "7. REC Team Dashboard",
        "8. Certificate Generator",
        "9. BI Dashboard and Reporting",
        "10. Documents, logs and monitoring",
        "11. Status guide and troubleshooting",
        "12. Daily Admin checklist",
    )
    table = doc.add_table(rows=0, cols=1)
    table.style = "Light Shading Accent 1"
    for item in sections:
        table.add_row().cells[0].text = item
    doc.add_paragraph()
    add_callout(
        doc,
        "Screenshot privacy",
        "Names, email addresses, student numbers and supervisor details are blurred in this manual. The live system displays authorised records normally.",
        LIGHT_GREY,
    )
    doc.add_page_break()


def build_manual():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc)

    doc.add_heading("1. Access and navigation", level=1)
    doc.add_heading("1.1 Sign in", level=2)
    add_steps(doc, (
        "Open the Integrated Ethics and MBA application in a supported browser.",
        "Select the Ethics system, then enter your authorised Admin account details.",
        "After successful authentication, confirm that the orange JBS Ethics Application sidebar is visible.",
        "Select Dashboard to begin processing submissions.",
    ))
    add_callout(doc, "Access denied", "If an Admin menu is missing, confirm that the account is active and has the ADMIN or SUPER_ADMIN role. Do not share Admin credentials.")
    doc.add_heading("1.2 Main Admin menu", level=2)
    menu = doc.add_table(rows=1, cols=2)
    menu.style = "Table Grid"
    menu.alignment = WD_TABLE_ALIGNMENT.CENTER
    menu.rows[0].cells[0].text = "Menu item"
    menu.rows[0].cells[1].text = "Purpose"
    for cell in menu.rows[0].cells:
        shade(cell, DARK_BLUE)
        set_cell_text_color(cell, WHITE)
    set_repeat_table_header(menu.rows[0])
    entries = (
        ("Dashboard", "Open submission folders by form type, year and month."),
        ("Reassign Supervisors", "Change a supervisor only while the application is with the student."),
        ("Reassign Reviewers", "Manage reviewer assignments for eligible applications."),
        ("Student Password Resets", "Process student password reset requests."),
        ("User Management", "Search, filter, activate, edit and manage users and student forms."),
        ("Upload Student Documents", "Upload authorised supporting documents."),
        ("Login Logs", "Review sign-in activity."),
        ("Admin Status Monitor", "Monitor application progress and operational status."),
        ("Certificate Generator", "Prepare, preview, edit, issue and send ethics certificates."),
        ("BI Dashboard and Reporting", "Configure and view authorised BI reports."),
        ("REC Team Dashboard", "Track applications submitted to the REC workflow."),
    )
    for left, right in entries:
        cells = menu.add_row().cells
        cells[0].text, cells[1].text = left, right
    doc.add_page_break()

    doc.add_heading("2. Admin dashboard and submission folders", level=1)
    add_screenshot(doc, "01_admin_dashboard.png", "Figure 1 — Admin Dashboard and submission folders")
    doc.add_heading("2.1 Find a submission", level=2)
    add_steps(doc, (
        "Select the required year at the top of the dashboard.",
        "Locate Form A, Form B or Form C.",
        "Find the required submission month and select View.",
        "Open the student record and verify the form status, risk level, documents and review history before acting.",
    ))
    doc.add_heading("2.2 Export dashboard information", level=2)
    add_bullets(doc, (
        "Export Forms (Excel) exports application information.",
        "Export Agenda (Excel) creates an agenda-oriented export.",
        "Export Reviewer Assignments (Excel) exports reviewer allocation information.",
    ))
    add_callout(doc, "Good practice", "Always verify the form type and current status before making a routing decision. A status change affects which user can act next.")

    doc.add_heading("3. Reviewing and routing an application", level=1)
    doc.add_heading("3.1 Review the submission", level=2)
    add_steps(doc, (
        "Open the application from its submission folder or the Certificate Generator.",
        "Review the applicant answers, uploaded documents, supervisor decision and reviewer feedback.",
        "Confirm the risk level and the latest reviewer decision.",
        "Choose only an action that is valid for the current status.",
    ))
    doc.add_heading("3.2 Decision guide", level=2)
    decision = doc.add_table(rows=1, cols=3)
    decision.style = "Table Grid"
    for index, text in enumerate(("Reviewer outcome", "Admin options", "Recommended check")):
        decision.rows[0].cells[index].text = text
        shade(decision.rows[0].cells[index], DARK_BLUE)
        set_cell_text_color(decision.rows[0].cells[index], WHITE)
    set_repeat_table_header(decision.rows[0])
    rows = (
        ("Approved", "Generate a certificate or send to REC, according to the Admin decision and institutional process.", "Confirm risk level and complete reviewer feedback."),
        ("Approved with minor changes", "Return to the student for resubmission, generate a certificate where appropriate, or send to REC.", "Confirm whether the changes must be completed before approval."),
        ("Medium/high risk or REC required", "Send the application to REC.", "Check supporting documents and risk classification."),
        ("With student", "Wait for resubmission; supervisor reassignment is permitted only in this state.", "Avoid routing stale or incomplete data."),
    )
    for values in rows:
        cells = decision.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    add_callout(doc, "Duplicate reviews", "A supervisor cannot review again when the status is With Ethics, and a reviewer cannot review again when the status is Reviewed. The system displays an already-reviewed message.")
    doc.add_page_break()

    doc.add_heading("4. User management", level=1)
    add_screenshot(doc, "02_user_management.png", "Figure 2 — User Management with search and role filters")
    doc.add_heading("4.1 Find and filter users", level=2)
    add_steps(doc, (
        "Select User Management from the sidebar.",
        "Enter a name, email address or role in the search box if required.",
        "Select an authentication status and/or role from the filters.",
        "Select Search to apply the filters.",
    ))
    doc.add_heading("4.2 Common actions", level=2)
    add_bullets(doc, (
        "Add User creates a standard account; Add Super Admin is restricted to authorised administration.",
        "Activate or Deactivate controls whether a user may access the system.",
        "Edit updates authorised account information and role settings.",
        "Delete should be used only under the approved data-governance procedure.",
        "When the STUDENT role filter is applied, use the student forms action to open and manage that student's forms.",
        "Archived Forms opens the archive register and restore workflow.",
    ))
    add_callout(doc, "Account checks", "Before creating a student account, search both active and inactive users using the full student email address. If an account appears to exist incorrectly, check alternate roles and authentication statuses.")

    doc.add_heading("5. Supervisors, reviewers and password resets", level=1)
    doc.add_heading("5.1 Reassign a supervisor", level=2)
    add_steps(doc, (
        "Open Reassign Supervisors and locate the student.",
        "Confirm that the application is currently with the student.",
        "Select the replacement supervisor and submit the change.",
        "Confirm the success message and verify the application record.",
    ))
    add_callout(doc, "Important restriction", "A supervisor cannot be changed after the student has submitted the form. Wait until the application is returned to the student before reassigning the supervisor.")
    doc.add_heading("5.2 Reassign a reviewer", level=2)
    add_steps(doc, (
        "Open Reassign Reviewers and locate the eligible application.",
        "Review the existing assignments and select the appropriate reviewer.",
        "Save the assignment, then confirm it appears in the application history or export.",
    ))
    doc.add_heading("5.3 Student password reset", level=2)
    add_steps(doc, (
        "Open Student Password Resets.",
        "Verify the student's identity using the approved internal process.",
        "Complete the reset action and instruct the student to sign in securely.",
    ))

    doc.add_heading("6. Archived student forms", level=1)
    doc.add_heading("6.1 Move forms to the archive", level=2)
    add_steps(doc, (
        "In User Management, filter by the STUDENT role.",
        "Open the student's forms page.",
        "Select one form, several forms, or Select All.",
        "Choose the archive action and confirm the selection.",
        "Verify that the record appears under Archived Forms with the name of the user who archived it.",
    ))
    doc.add_heading("6.2 Restore an archived form", level=2)
    add_steps(doc, (
        "Open User Management, then select Archived Forms.",
        "Locate the correct student and archived form.",
        "Select Restore.",
        "Confirm that the archived form type matches any new form already started by the student.",
        "After restoration, ask the student to sign in and verify that the form appears under Form A, B or C and continues from the saved point.",
    ))
    add_callout(doc, "Form-type protection", "A Form B archive cannot replace a newly started Form A or Form C. Matching the form type prevents answers and workflow data from being placed into the wrong application structure. If the types differ, cancel the restore and resolve the active form first.")
    doc.add_page_break()

    doc.add_heading("7. REC Team Dashboard", level=1)
    add_screenshot(doc, "04_rec_dashboard.png", "Figure 3 — REC Team Dashboard")
    doc.add_heading("7.1 Track REC submissions", level=2)
    add_steps(doc, (
        "Select REC Team Dashboard.",
        "Locate the application by form type, risk level, date assigned and status.",
        "Use Check Feedback to open the REC review details.",
        "Verify the review count and supporting documents before recording or communicating an outcome.",
    ))
    add_callout(doc, "Missing REC form", "Confirm that the application was actually submitted to REC, then check the correct form type and status. Refresh the dashboard after the transaction completes.")

    doc.add_heading("8. Certificate Generator", level=1)
    add_screenshot(doc, "03_certificate_generator.png", "Figure 4 — Certificate Generator with form-type and status filters")
    doc.add_heading("8.1 Find an eligible application", level=2)
    add_steps(doc, (
        "Select Certificate Generator.",
        "Choose Form A, Form B or Form C from Form type.",
        "Use Filter Forms by Status to narrow the list.",
        "Check the review status, risk level, reviewer feedback and supporting documents.",
        "Select Review for the required application.",
    ))
    doc.add_heading("8.2 Prepare and issue a certificate", level=2)
    add_steps(doc, (
        "Start or continue the certificate draft for the eligible application.",
        "Review and edit the certificate fields carefully. Save the draft before leaving the page.",
        "Preview the certificate and check names, student number, project title, approval wording, dates and conditions.",
        "Issue the certificate only when all details are correct and approval is authorised.",
        "Use the send option to deliver the issued certificate through the configured workflow.",
        "Reopen the record to confirm the issue status and certificate reference.",
    ))
    add_callout(doc, "Approved applications", "An Approved reviewer decision may still present both Generate Certificate and Send to REC options so the Admin can make the final authorised routing decision.")
    add_callout(doc, "Safe editing", "Edit the saved draft and use Preview before issuing. Treat issuance as the formal approval event; do not issue a certificate merely to inspect its layout.", LIGHT_GREY)

    doc.add_heading("9. BI Dashboard and Reporting", level=1)
    add_screenshot(doc, "05_bi_reporting.png", "Figure 5 — BI Dashboard and Reporting")
    doc.add_heading("9.1 View a BI report", level=2)
    add_steps(doc, (
        "Select BI Dashboard and Reporting.",
        "Use Search, BI View Name, Database Table, Status and Records Per Page to filter the configured views.",
        "Select Apply.",
        "Select View BI Report for the required active view.",
    ))
    doc.add_heading("9.2 Configure BI rights and templates", level=2)
    add_steps(doc, (
        "Select Configuration Rights if your account is authorised to administer BI access.",
        "Grant only the minimum rights required for each role or user and save the configuration.",
        "Return to the BI dashboard and verify the saved configuration.",
        "Use Configure BI Template to select and arrange the permitted report columns, then save and preview the report.",
    ))
    add_callout(doc, "If rights are not saved", "Confirm that the user and report selection are valid, submit the form once, and check for an on-screen validation message. If the problem continues, capture the time, user, report and server error for technical support.")
    doc.add_page_break()

    doc.add_heading("10. Documents, logs and monitoring", level=1)
    doc.add_heading("10.1 Upload student documents", level=2)
    add_steps(doc, (
        "Open Upload Student Documents.",
        "Select the correct student, form and document category.",
        "Choose the approved file and upload it.",
        "Open the application afterward to confirm the document appears in the correct section.",
    ))
    doc.add_heading("10.2 Login Logs", level=2)
    add_bullets(doc, (
        "Use Login Logs to investigate access and sign-in questions.",
        "Match the account, timestamp and action before drawing a conclusion.",
        "Handle exported or copied audit information as confidential data.",
    ))
    doc.add_heading("10.3 Admin Status Monitor", level=2)
    add_bullets(doc, (
        "Use the monitor to identify applications waiting at each workflow stage.",
        "Open the underlying application before changing or communicating a status.",
        "The former super-admin monitoring pages are not part of the supported workflow; use Admin Status Monitor and BI reporting.",
    ))

    doc.add_heading("11. Status guide and troubleshooting", level=1)
    status = doc.add_table(rows=1, cols=3)
    status.style = "Table Grid"
    for i, label in enumerate(("Status/issue", "Meaning or likely cause", "Admin action")):
        status.rows[0].cells[i].text = label
        shade(status.rows[0].cells[i], DARK_BLUE)
        set_cell_text_color(status.rows[0].cells[i], WHITE)
    set_repeat_table_header(status.rows[0])
    issues = (
        ("With student", "The student can edit or resubmit.", "Supervisor may be reassigned if required."),
        ("With ethics", "The supervisor stage is already complete.", "Do not permit another supervisor review."),
        ("Reviewed", "Reviewer action is complete.", "Do not permit duplicate reviewer review; assess routing options."),
        ("Submitted to REC", "The application has entered the REC workflow.", "Use REC Team Dashboard and check feedback."),
        ("Approved", "Reviewer/committee outcome permits an Admin decision.", "Generate a certificate or send to REC as authorised."),
        ("Internal Server Error", "A server, database, route or validation error occurred.", "Record URL, time, action and visible message; check the application log before retrying."),
        ("Access denied", "Role, activation or session does not permit the page.", "Verify the active account role and sign in again."),
        ("Restored form not visible", "Form type/status or restoration record may conflict.", "Confirm matching form type and that restore completed; do not start another form."),
    )
    for values in issues:
        cells = status.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value

    doc.add_heading("12. Daily Admin checklist", level=1)
    add_bullets(doc, (
        "Check submission folders and applications waiting for Admin action.",
        "Review supervisor and reviewer assignment queues.",
        "Check applications submitted to REC and outstanding feedback.",
        "Process eligible certificate drafts, previews and authorised issuance.",
        "Review password reset and account activation requests.",
        "Check status monitoring and BI reports for stalled records.",
        "Record and escalate errors with the URL, timestamp, user role and action performed.",
        "Log out when administration work is complete.",
    ))
    add_callout(doc, "Final principle", "Check the current form type, status and latest decision before every administrative action. This protects the student's saved work and keeps the audit trail consistent.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    build_manual()
